"""Compact DOCX export matching the print-ready PDF exam structure."""
from __future__ import annotations

import base64
import binascii
import io
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from app.exports.common import (
    HEADER_LABELS_BY_LANG,
    RESPONSE_LINE_TEXT,
    TRUE_FALSE_CHOICES_BY_LANG,
    group_exam_sections,
    response_line_count,
)

PAGE_MARGIN_IN = 0.5
CONTENT_WIDTH_IN = 7.27  # A4 width (8.27in) minus two 0.5in margins
FONT_NAME = "Arial"
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(100, 116, 139)
SECTION_FILL = "E8EEF5"
RULE = "B4BECA"


def _labels(language: str) -> dict[str, str]:
    return HEADER_LABELS_BY_LANG.get(language, HEADER_LABELS_BY_LANG["en"])


def _set_paragraph_rtl(paragraph) -> None:
    """Mark a paragraph as bidi so Word lays it out right-to-left."""
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def _clean(value: object, fallback: str = "") -> str:
    text = str(value or "").strip()
    return text or fallback


def _image_stream(data_url: object) -> io.BytesIO | None:
    if not isinstance(data_url, str) or not data_url.startswith("data:image/"):
        return None
    try:
        _, encoded = data_url.split(",", 1)
        raw = base64.b64decode(encoded, validate=True)
    except (ValueError, binascii.Error):
        return None
    if not raw or len(raw) > 2_500_000:
        return None
    return io.BytesIO(raw)


def _set_run_font(
    run,
    *,
    size: float = 10.5,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor = INK,
) -> None:
    run.font.name = FONT_NAME
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    # Complex-script font so Arabic glyphs use the same face and size.
    r_fonts.set(qn("w:cs"), FONT_NAME)
    cs_size = r_pr.find(qn("w:szCs"))
    if cs_size is None:
        cs_size = OxmlElement("w:szCs")
        r_pr.append(cs_size)
    cs_size.set(qn("w:val"), str(int(size * 2)))
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
        b_cs = OxmlElement("w:bCs")
        r_pr.append(b_cs)
    if italic is not None:
        run.italic = italic


def _shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _paragraph_bottom_border(paragraph, color: str = RULE, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def _set_cell_width(cell, width_in: float) -> None:
    width_dxa = str(round(width_in * 1440))
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), width_dxa)
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_in)


def _set_cell_margins(cell, *, top: int = 60, start: int = 80, bottom: int = 60, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_in: list[float], *, indent_dxa: int = 0) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_width_dxa = round(sum(widths_in) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width in zip(row.cells, widths_in):
            _set_cell_width(cell, width)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.05

    heading = doc.styles["Heading 1"]
    heading.font.name = FONT_NAME
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    heading.font.size = Pt(11)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(51, 65, 85)
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(5)
    heading.paragraph_format.keep_with_next = True


def _add_field(paragraph, instruction: str, placeholder: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    _set_run_font(run, size=8, color=MUTED)


def _configure_section(section, title: str, model_number: int, copy_label: str) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(PAGE_MARGIN_IN)
    section.right_margin = Inches(PAGE_MARGIN_IN)
    section.bottom_margin = Inches(PAGE_MARGIN_IN)
    section.left_margin = Inches(PAGE_MARGIN_IN)
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.18)

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(0)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run("Page ")
    _set_run_font(run, size=8, color=MUTED)
    _add_field(fp, "PAGE", "1")
    run = fp.add_run(" of ")
    _set_run_font(run, size=8, color=MUTED)
    _add_field(fp, "NUMPAGES", "1")


def _add_logo(cell, data_url: object, alignment: WD_ALIGN_PARAGRAPH) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(0)
    stream = _image_stream(data_url)
    if stream is None:
        return
    try:
        p.add_run().add_picture(stream, width=Inches(0.72), height=Inches(0.46))
    except Exception:
        return


def _add_exam_header(doc: Document, metadata: dict[str, Any], model_number: int, language: str = "en") -> None:
    rtl = language == "ar"
    labels = _labels(language)
    title = _clean(metadata.get("exam_title"), "Examination")
    table = doc.add_table(rows=1, cols=3)
    _set_table_geometry(table, [1.0, CONTENT_WIDTH_IN - 2.0, 1.0])
    _remove_table_borders(table)
    left_align = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    right_align = WD_ALIGN_PARAGRAPH.LEFT if rtl else WD_ALIGN_PARAGRAPH.RIGHT
    _add_logo(table.cell(0, 0), metadata.get("left_logo_data"), left_align)
    _add_logo(table.cell(0, 2), metadata.get("right_logo_data"), right_align)
    center = table.cell(0, 1)
    center.text = ""
    p = center.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    _set_run_font(run, size=16, bold=True)
    p2 = center.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(labels["model"].format(n=model_number))
    _set_run_font(run, size=9, color=MUTED)

    fields = [
        (labels["class"], _clean(metadata.get("class_name"), "_________")),
        (labels["duration"], _clean(metadata.get("duration"), "_______")),
        (labels["date"], _clean(metadata.get("exam_date"), "_________")),
        (labels["teacher"], _clean(metadata.get("teacher_name"), "_________")),
    ]
    meta = doc.add_table(rows=1, cols=4)
    widths = [CONTENT_WIDTH_IN / 4] * 4
    _set_table_geometry(meta, widths)
    for cell, (label, value) in zip(meta.rows[0].cells, fields):
        cell.text = ""
        para = cell.paragraphs[0]
        _shade_paragraph(para, "F6F8FB")
        if rtl:
            _set_paragraph_rtl(para)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"{label}: ")
        _set_run_font(run, size=8, bold=True, color=MUTED)
        run = para.add_run(value)
        _set_run_font(run, size=8.5)

    student = doc.add_paragraph()
    student.paragraph_format.space_before = Pt(3)
    student.paragraph_format.space_after = Pt(5)
    if rtl:
        _set_paragraph_rtl(student)
        student.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = student.add_run("_" * 38)
        _set_run_font(run, size=9.5)
        run = student.add_run("        " + labels["student_name"])
        _set_run_font(run, size=9.5)
        run = student.add_run("_" * 14)
        _set_run_font(run, size=9.5)
        run = student.add_run("        " + labels["class_suffix"])
        _set_run_font(run, size=9.5)
    else:
        run = student.add_run(labels["student_name"])
        _set_run_font(run, size=9.5)
        run = student.add_run("_" * 38)
        _set_run_font(run, size=9.5)
        run = student.add_run("        " + labels["class_suffix"])
        _set_run_font(run, size=9.5)
        run = student.add_run("_" * 14)
        _set_run_font(run, size=9.5)


def _add_section_heading(doc: Document, label: str, *, page_break_before: bool = False, rtl: bool = False) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.left_indent = Pt(4)
    p.paragraph_format.right_indent = Pt(4)
    if rtl:
        _set_paragraph_rtl(p)
    p.add_run(label)
    _shade_paragraph(p, SECTION_FILL)
    _paragraph_bottom_border(p, color=RULE, size="5")


def _add_question_stem(doc: Document, item: dict[str, Any], rtl: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    if rtl:
        _set_paragraph_rtl(p)
    run = p.add_run(f"{item['number']}. " if rtl else f"Q{item['number']}. ")
    _set_run_font(run, bold=True)
    run = p.add_run(_clean(item.get("text"), "(missing question text)"))
    _set_run_font(run, bold=True)


def _add_mcq(doc: Document, item: dict[str, Any], rtl: bool = False) -> None:
    _add_question_stem(doc, item, rtl=rtl)
    options = list((item.get("options") or {}).items())
    two_column = options and all(len(_clean(option)) <= 90 for _, option in options)
    if two_column:
        rows = (len(options) + 1) // 2
        table = doc.add_table(rows=rows, cols=2)
        _set_table_geometry(table, [CONTENT_WIDTH_IN / 2] * 2)
        _remove_table_borders(table)
        for index, (letter, option) in enumerate(options):
            cell = table.cell(index // 2, index % 2)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if rtl:
                _set_paragraph_rtl(p)
            run = p.add_run(f"{_clean(option)} .{letter}" if rtl else f"{letter}. {_clean(option)}")
            _set_run_font(run, size=9.5)
    else:
        for letter, option in options:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(1)
            if rtl:
                _set_paragraph_rtl(p)
            run = p.add_run(f"{_clean(option)} .{letter}" if rtl else f"{letter}. {_clean(option)}")
            _set_run_font(run, size=9.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_word_bank(doc: Document, words: list[str], language: str = "en") -> None:
    if not words:
        return
    from app.online.models import UI_STRINGS_BY_LANG

    strings = UI_STRINGS_BY_LANG.get(language, UI_STRINGS_BY_LANG["en"])
    table = doc.add_table(rows=1, cols=1)
    _set_table_geometry(table, [CONTENT_WIDTH_IN])
    cell = table.cell(0, 0)
    cell.text = ""
    _shade_paragraph(cell.paragraphs[0], "F8FAFC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(strings["word_bank"].upper() if language == "en" else strings["word_bank"])
    _set_run_font(run, size=8.5, bold=True, color=MUTED)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run("   |   ".join(words))
    _set_run_font(run, size=9.5)


def _add_answer_lines(doc: Document, count: int) -> None:
    for index in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0 if index < count - 1 else 3)
        p.paragraph_format.line_spacing = Pt(16)
        # Explicit underscore glyphs are used instead of paragraph borders.
        # Some Word/viewer combinations suppress borders on otherwise empty
        # paragraphs, which made the response area look like blank spacing.
        run = p.add_run(RESPONSE_LINE_TEXT)
        _set_run_font(run, size=9, color=RGBColor(71, 85, 105))


def _answer_text(item: dict[str, Any], language: str = "en") -> str:
    qtype = item["qtype"]
    if qtype == "mcq":
        return _clean(item.get("correct_answer"), "No answer supplied")
    if qtype == "true_false":
        from app.online.models import tf_answer_label

        raw = _clean(item.get("answer"), "No answer supplied")
        if raw.lower() in {"true", "false"}:
            return tf_answer_label(raw, language)
        return raw
    if qtype == "fill_in_the_blank":
        return ", ".join(_clean(answer) for answer in (item.get("answers") or [])) or "No answer supplied"
    answer = _clean(item.get("reference_answer"), "No reference answer supplied")
    points = [_clean(point) for point in (item.get("key_points") or []) if _clean(point)]
    return answer + ((" | Key points: " + "; ".join(points)) if points else "")


def _render_student_sections(doc: Document, sections: list[dict[str, Any]], language: str = "en") -> None:
    rtl = language == "ar"
    tf_choices = TRUE_FALSE_CHOICES_BY_LANG.get(language, TRUE_FALSE_CHOICES_BY_LANG["en"])
    for section_index, section in enumerate(sections):
        _add_section_heading(
            doc,
            section["label"],
            page_break_before=section["qtype"] == "essay" and section_index > 0,
            rtl=rtl,
        )
        if section["qtype"] == "fill_in_the_blank":
            _add_word_bank(doc, section.get("word_bank") or [], language=language)
        for item in section["items"]:
            qtype = section["qtype"]
            if qtype == "mcq":
                _add_mcq(doc, item, rtl=rtl)
            elif qtype == "true_false":
                statement = _clean(item.get("text"))
                choices_below = len(statement) > 80
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = choices_below
                if rtl:
                    _set_paragraph_rtl(p)
                run = p.add_run(f"{item['number']}. {statement}" if rtl else f"Q{item['number']}. {statement}")
                _set_run_font(run, bold=True)
                if choices_below:
                    choices = doc.add_paragraph()
                    choices.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    choices.paragraph_format.space_after = Pt(4)
                    run = choices.add_run(tf_choices)
                    _set_run_font(run, size=9.5)
                else:
                    run = p.add_run("     " + tf_choices)
                    _set_run_font(run, size=9.5)
            elif qtype == "fill_in_the_blank":
                _add_question_stem(doc, item, rtl=rtl)
            else:
                _add_question_stem(doc, item, rtl=rtl)
                _add_answer_lines(doc, response_line_count(qtype))


def _render_answer_key(
    doc: Document,
    sections: list[dict[str, Any]],
    model_number: int,
    title: str,
    language: str = "en",
) -> None:
    rtl = language == "ar"
    labels = _labels(language)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(section, title, model_number, "Teacher answer key")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(labels["answer_key"].format(n=model_number))
    _set_run_font(run, size=15, bold=True)
    for section in sections:
        _add_section_heading(doc, section["label"], rtl=rtl)
        for item in section["items"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            if rtl:
                _set_paragraph_rtl(p)
            run = p.add_run(f"{item['number']}. " if rtl else f"Q{item['number']}. ")
            _set_run_font(run, size=9.5, bold=True)
            run = p.add_run(_answer_text(item, language=language))
            _set_run_font(run, size=9.5)


def render_exam_docx(stored_record: dict[str, Any]) -> bytes:
    """Render all stored models as compact A4 student exams + teacher keys."""
    doc = Document()
    _configure_styles(doc)
    metadata = dict(stored_record.get("metadata") or {})
    language = str(metadata.get("document_language") or "en")
    rendered_models = 0

    for exam in stored_record.get("exams") or []:
        if not exam.get("questions"):
            continue
        model_number = int(exam.get("model_number") or 1)
        title = _clean(metadata.get("exam_title") or exam.get("title"), "Examination")
        if rendered_models:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
        else:
            section = doc.sections[0]
        _configure_section(section, title, model_number, "Student copy")
        _add_exam_header(doc, metadata, model_number, language=language)
        sections = group_exam_sections(exam.get("questions") or {}, language=language)
        _render_student_sections(doc, sections, language=language)
        _render_answer_key(doc, sections, model_number, title, language=language)
        rendered_models += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
